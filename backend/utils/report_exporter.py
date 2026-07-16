"""
报告导出工具 — 分层报告树 + Markdown / Word / HTML

借鉴:
  - TradingAgents 的 write_report_tree(): 按 Agent 团队分目录保存
  - TradingAgents 的 render_*(): Pydantic → 确定性 Markdown 渲染
  - data_analysis_agent 的 LLM 叙事 + 相对路径图表引用

v3.3 改进:
  - 分层报告树: 按阶段保存 1_plan/ 2_sql/ 3_chart/ 4_debate/ 5_final/
  - 完整产物合并: complete_report.md 包含所有阶段的摘要
  - DOCX 增强: 使用 markdown AST 解析而非逐行正则
  - 报告元信息: 用户 query + 表名 + 时间 + 状态
  - 结构附录: SQL原文 + 辩论原文 + 评估结果
"""

import logging
from datetime import datetime
from pathlib import Path
from typing import Any

from backend.agent.schemas import (
    render_chart_section,
    render_debate_scores,
    render_performance,
    render_plan,
    render_sql_section,
    render_validation,
)

logger = logging.getLogger("utils")

EXPORT_DIR = Path("data/exports")


# ═══════════════════════════════════════════════════════════
# 公共入口
# ═══════════════════════════════════════════════════════════


def export_report(
    final_state: dict[str, Any],
    format: str = "md",
    output_dir: str = "data/exports",
) -> str | None:
    """
    导出分析报告（单文件）。

    Args:
        final_state: 分析完成后的最终状态
        format:      "md" | "docx" | "html"
        output_dir:  输出目录

    Returns:
        导出文件路径，失败返回 None
    """
    output_path = Path(output_dir)
    output_path.mkdir(parents=True, exist_ok=True)

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    query_slug = _slugify(str(final_state.get("user_query", "analysis"))[:30])
    filename = f"{timestamp}_{query_slug}"

    content = build_complete_report(final_state)

    try:
        if format == "md":
            return _export_md(content, output_path, filename)
        elif format == "docx":
            return _export_docx(
                content, output_path, filename,
                chart_files=final_state.get("chart_files", []),
                chart_json=final_state.get("chart_json"),
            )
        elif format == "html":
            return _export_html(
                content, output_path, filename,
                chart_json=final_state.get("chart_json"),
            )
        else:
            logger.warning("不支持的导出格式: %s", format)
            return None
    except Exception as e:
        logger.error("报告导出失败: %s", e)
        return None


def export_report_tree(
    final_state: dict[str, Any],
    output_dir: str = "data/exports",
) -> Path:
    """
    导出分层报告树（借鉴 TradingAgents 的 write_report_tree）。

    按阶段保存所有 Agent 产出到子目录:
      <output_dir>/<timestamp>_<query>/
      ├── 1_plan/
      │   └── plan.md
      ├── 2_sql/
      │   ├── query.sql
      │   └── result.csv
      ├── 3_chart/
      │   └── chart.json (或 PNG 引用)
      ├── 4_debate/
      │   ├── optimistic.md
      │   ├── pessimistic.md
      │   └── scores.md
      ├── 5_report/
      │   ├── draft.md
      │   └── validation.md
      └── complete_report.md

    Returns:
        报告目录 Path
    """
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    query_slug = _slugify(str(final_state.get("user_query", "analysis"))[:30])
    root = Path(output_dir) / f"{timestamp}_{query_slug}"
    root.mkdir(parents=True, exist_ok=True)

    # ─── 1. 计划 ───
    plan = final_state.get("plan", [])
    if plan:
        plan_dir = root / "1_plan"
        plan_dir.mkdir(exist_ok=True)
        (plan_dir / "plan.md").write_text(
            f"# 执行计划\n\n{render_plan(plan)}", encoding="utf-8"
        )

    # ─── 2. SQL ───
    sql_query = final_state.get("sql_query", "")
    sql_result = final_state.get("sql_result", "")
    if sql_query or sql_result:
        sql_dir = root / "2_sql"
        sql_dir.mkdir(exist_ok=True)
        if sql_query:
            (sql_dir / "query.sql").write_text(sql_query, encoding="utf-8")
        if sql_result:
            (sql_dir / "result.csv").write_text(sql_result, encoding="utf-8")

    # ─── 3. 图表 ───
    chart_json = final_state.get("chart_json")
    chart_files = final_state.get("chart_files", [])
    if chart_json or chart_files:
        import shutil
        chart_dir = root / "3_chart"
        chart_dir.mkdir(exist_ok=True)

        # 复制 PNG 文件到报告树目录
        png_copied = []
        for cf in chart_files:
            src = Path(cf)
            if src.exists():
                dst = chart_dir / src.name
                shutil.copy2(src, dst)
                png_copied.append(str(dst))
                logger.info("[Export] 图表已复制: %s -> %s", src.name, dst.name)
            else:
                logger.warning("[Export] 图表文件不存在: %s", cf)

        chart_section = render_chart_section(chart_json, png_copied)
        if chart_section:
            (chart_dir / "chart.md").write_text(f"# 图表\n\n{chart_section}", encoding="utf-8")
        if chart_json and isinstance(chart_json, dict):
            import json
            (chart_dir / "chart.json").write_text(
                json.dumps(chart_json, ensure_ascii=False, indent=2), encoding="utf-8"
            )

    # ─── 4. 辩论 ───
    optimistic = final_state.get("optimistic_view", "")
    pessimistic = final_state.get("pessimistic_view", "")
    debate_scores = final_state.get("debate_scores")
    if optimistic or pessimistic:
        debate_dir = root / "4_debate"
        debate_dir.mkdir(exist_ok=True)
        if optimistic:
            (debate_dir / "optimistic.md").write_text(
                f"# 正方（乐观方）\n\n{optimistic}", encoding="utf-8"
            )
        if pessimistic:
            (debate_dir / "pessimistic.md").write_text(
                f"# 反方（谨慎方）\n\n{pessimistic}", encoding="utf-8"
            )
        if debate_scores:
            (debate_dir / "scores.md").write_text(
                f"# 辩论评分\n\n{render_debate_scores(debate_scores)}", encoding="utf-8"
            )

    # ─── 5. 报告与验证 ───
    report = final_state.get("final_report", "") or final_state.get("draft_report", "")
    validation = final_state.get("validation_result", "")
    report_dir = root / "5_report"
    report_dir.mkdir(exist_ok=True)
    if report:
        (report_dir / "draft.md").write_text(report, encoding="utf-8")
    if validation:
        val_dict = {
            "result": validation,
            "reason": final_state.get("validation_reason", ""),
            "revise_suggestions": final_state.get("validation_revise_suggestions", ""),
        }
        (report_dir / "validation.md").write_text(
            f"# 验证结果\n\n{render_validation(val_dict)}", encoding="utf-8"
        )

    # ─── 完整报告 ───
    complete = build_complete_report(final_state)
    (root / "complete_report.md").write_text(complete, encoding="utf-8")

    logger.info("[Export] 分层报告树已保存: %s (%d 个子目录)", root,
                len([d for d in root.iterdir() if d.is_dir()]))
    return root


# ═══════════════════════════════════════════════════════════
# 完整报告构建（借鉴 TradingAgents + data_analysis_agent）
# ═══════════════════════════════════════════════════════════


def build_complete_report(final_state: dict[str, Any]) -> str:
    """
    构建完整分析报告（Markdown）。

    结构:
      - 报告头部（用户 query + 表名 + 时间 + 状态）
      - I. 执行计划
      - II. 数据分析（LLM 叙事）
      - III. 图表
      - IV. 辩论与验证
      - V. 最终结论
      - 附录 A: SQL 原文
      - 附录 B: 辩论原文
      - 附录 C: 性能统计

    借鉴:
      - data_analysis_agent: LLM 叙事为主，相对路径图表引用
      - TradingAgents: 分阶段 append，Pydantic render 保证格式稳定性
    """
    sections: list[str] = []

    # ─── 报告头部 ───
    header = _build_header(final_state)
    sections.append(header)

    # ─── I. 执行计划 ───
    plan = final_state.get("plan", [])
    if plan:
        sections.append(f"## I. 执行计划\n\n{render_plan(plan)}\n")

    # ─── II. 数据分析（LLM 叙事） ───
    report = final_state.get("final_report", "") or final_state.get("draft_report", "")
    if report:
        sections.append(f"## II. 数据分析\n\n{report}\n")

    # ─── III. 图表 ───
    chart_json = final_state.get("chart_json")
    chart_files = final_state.get("chart_files", [])
    # Markdown 版本（base64 PNG 降级如果有 kaleido，否则提示文字）
    chart_section = render_chart_section(chart_json, chart_files, for_html=False)
    if chart_section:
        sections.append(f"## III. 数据可视化\n\n{chart_section}\n")

    # ─── IV. 辩论与验证（完整模式） ───
    debate_scores = final_state.get("debate_scores")
    validation = final_state.get("validation_result", "")
    if debate_scores or validation:
        sections.append("## IV. 质量评估\n")
        if debate_scores:
            sections.append(render_debate_scores(debate_scores))
            sections.append("")
        if validation:
            val_dict = {
                "result": validation,
                "reason": final_state.get("validation_reason", ""),
                "revise_suggestions": final_state.get("validation_revise_suggestions", ""),
            }
            sections.append(f"### 验证结果\n\n{render_validation(val_dict)}\n")

    # ─── V. 最终结论 ───
    if report:
        sections.append("## V. 最终结论\n")
        # 从报告中提取最后一段作为结论，或保留完整报告的最后部分
        sections.append("*详见「II. 数据分析」部分的结论与建议章节*\n")

    # ═══ 附录 ═══
    sections.append("---\n")
    sections.append("## 附录\n")

    # 附录 A: SQL
    sql_query = final_state.get("sql_query", "")
    sql_result = final_state.get("sql_result", "")
    if sql_query or sql_result:
        sql_section = render_sql_section(sql_query, sql_result)
        sections.append(f"### 附录 A: SQL 查询\n\n{sql_section}\n")

    # 附录 B: 辩论原文（完整模式）
    optimistic = final_state.get("optimistic_view", "")
    pessimistic = final_state.get("pessimistic_view", "")
    if optimistic or pessimistic:
        sections.append("### 附录 B: 辩论原文\n")
        if optimistic:
            sections.append(f"<details>\n<summary>🔵 正方（乐观方）</summary>\n\n{optimistic}\n\n</details>\n")
        if pessimistic:
            sections.append(f"<details>\n<summary>🔴 反方（谨慎方）</summary>\n\n{pessimistic}\n\n</details>\n")

    # 附录 C: 性能统计
    performance = final_state.get("performance_metrics")
    if performance:
        sections.append(f"### 附录 C: 性能统计\n\n{render_performance(performance)}\n")

    # ─── 页脚 ───
    sections.append("---\n")
    sections.append(
        f"*本报告由 DataForge AI 自动生成 | 7 Agents + LangGraph | "
        f"{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}*"
    )

    return "\n\n".join(sections)


def _build_header(final_state: dict[str, Any]) -> str:
    """构建报告头部元信息。"""
    user_query = final_state.get("user_query", "未知查询")
    available_tables = final_state.get("available_tables", [])
    validation = final_state.get("validation_result", "")
    total_time = final_state.get("performance_metrics", {}).get("total_time", 0)

    status_emoji = {"approved": "✅", "rejected": "❌", "needs_review": "⚠️"}.get(validation, "📋")
    status_label = {"approved": "已通过验证", "rejected": "已驳回", "needs_review": "需人工审核"}.get(
        validation, "未验证"
    )

    lines = [
        "# DataForge AI 分析报告",
        "",
        f"| 项目 | 内容 |",
        f"|------|------|",
        f"| **分析问题** | {user_query[:100]} |",
        f"| **数据表** | {', '.join(available_tables) if available_tables else '—'} |",
        f"| **生成时间** | {datetime.now().strftime('%Y-%m-%d %H:%M:%S')} |",
        f"| **分析状态** | {status_emoji} {status_label} |",
        f"| **总耗时** | {total_time:.1f} 秒 |",
        "",
        "---",
    ]
    return "\n".join(lines)


# ═══════════════════════════════════════════════════════════
# 格式导出
# ═══════════════════════════════════════════════════════════


def _export_md(content: str, output_dir: Path, filename: str) -> str:
    """导出 Markdown。"""
    filepath = output_dir / f"{filename}.md"
    filepath.write_text(content, encoding="utf-8")
    logger.info("Markdown 报告已导出: %s", filepath)
    return str(filepath)


def _export_docx(
    content: str,
    output_dir: Path,
    filename: str,
    chart_files: list[str] | None = None,
    chart_json: dict | None = None,
) -> str:
    """
    导出 Word 文档 — base64 图片解码为真实图片嵌入。

    特性:
      - 提取 report 中的 data:image/png;base64,... 并嵌入为图片
      - chart_files 中的 PNG 也嵌入
      - markdown AST → 结构化 Word（表格/代码块/列表）
    """
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Inches, Pt
    except ImportError:
        logger.warning("python-docx 未安装，回退到 Markdown")
        return _export_md(content, output_dir, filename)

    import re as _re
    import base64 as _b64
    import io as _io

    doc = Document()

    # 提取并嵌入 base64 图片
    base64_imgs = list(_re.finditer(r'!\[([^\]]*)\]\((data:image/png;base64,([^)]+))\)', content))
    if base64_imgs:
        doc.add_heading("数据可视化", level=2)
        for m in base64_imgs:
            alt_text = m.group(1)
            b64_data = m.group(3)
            try:
                img_bytes = _b64.b64decode(b64_data)
                img_stream = _io.BytesIO(img_bytes)
                doc.add_picture(img_stream, width=Inches(5.5))
                last_paragraph = doc.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                if alt_text:
                    doc.add_paragraph(alt_text).alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception as e:
                logger.warning("DOCX base64 图片嵌入失败: %s", e)
        # 从 content 中移除 base64 图片引用（避免显示为文本）
        content = _re.sub(r'!\[[^\]]*\]\(data:image/png;base64,[^)]+\)', '', content)

    # ─── 尝试 markdown AST 解析 ───
    try:
        import markdown

        html = markdown.markdown(content, extensions=["tables", "fenced_code", "codehilite"])
        _html_to_docx(html, doc)
    except ImportError:
        _lines_to_docx(content, doc)

    # ─── 嵌入图表图片 ───
    all_pngs = list(chart_files or [])

    # 如果只有 Plotly JSON 没有 PNG，尝试转换
    if not all_pngs and chart_json and isinstance(chart_json, dict) and "data" in chart_json:
        try:
            import plotly.io as pio
            import json
            import tempfile
            fig = pio.from_json(json.dumps(chart_json))
            tmp_png = Path(tempfile.gettempdir()) / f"_dataforge_chart_{filename}.png"
            fig.write_image(str(tmp_png), format="png", width=800, height=450, scale=1)
            if tmp_png.exists():
                all_pngs.append(str(tmp_png))
                logger.info("Plotly JSON → PNG 转换成功: %s", tmp_png)
        except Exception as e:
            logger.debug("Plotly JSON → PNG 转换失败（kaleido 可能未安装）: %s", e)

    if all_pngs:
        doc.add_heading("数据可视化", level=2)
        for png_path in all_pngs:
            try:
                p = Path(png_path)
                if p.exists():
                    doc.add_picture(str(p), width=Inches(5.5))
                    last_paragraph = doc.paragraphs[-1]
                    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    doc.add_paragraph(p.name).alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    doc.add_paragraph(f"[图表文件未找到: {p.name}]")
            except Exception as e:
                logger.warning("图表嵌入失败 (%s): %s", png_path, e)

    # ─── 保存 ───
    filepath = output_dir / f"{filename}.docx"
    doc.save(str(filepath))
    logger.info("Word 报告已导出: %s (图表=%d)", filepath, len(all_pngs))
    return str(filepath)


def _html_to_docx(html: str, doc):
    """将 HTML 字符串转换为 Word 文档元素（简化版）。"""
    from docx import Document
    from docx.shared import Pt
    import re

    # 移除 HTML 标签，提取纯文本段落
    # 处理标题：<h1> → Heading 1, <h2> → Heading 2
    # 处理表格：<table> → Word Table
    # 处理代码块：<pre><code> → 等宽字体段落

    # 简单分段处理
    # 替换常见标签为段落标记
    text = html
    text = re.sub(r'<h1[^>]*>(.*?)</h1>', r'[H1]\1[/H1]', text, flags=re.DOTALL)
    text = re.sub(r'<h2[^>]*>(.*?)</h2>', r'[H2]\1[/H2]', text, flags=re.DOTALL)
    text = re.sub(r'<h3[^>]*>(.*?)</h3>', r'[H3]\1[/H3]', text, flags=re.DOTALL)
    text = re.sub(r'<li>(.*?)</li>', r'[LI]\1[/LI]', text, flags=re.DOTALL)
    text = re.sub(r'<pre><code[^>]*>(.*?)</code></pre>', r'[CODE]\1[/CODE]', text, flags=re.DOTALL)
    text = re.sub(r'<p>(.*?)</p>', r'\1\n', text, flags=re.DOTALL)
    text = re.sub(r'<br\s*/?>', '\n', text)
    text = re.sub(r'<[^>]+>', '', text)  # 移除剩余 HTML 标签

    # 按段落处理
    for para_text in text.split('\n'):
        para_text = para_text.strip()
        if not para_text:
            continue

        if para_text.startswith('[H1]') and '[/H1]' in para_text:
            content_text = para_text.replace('[H1]', '').replace('[/H1]', '')
            doc.add_heading(content_text, level=1)
        elif para_text.startswith('[H2]') and '[/H2]' in para_text:
            content_text = para_text.replace('[H2]', '').replace('[/H2]', '')
            doc.add_heading(content_text, level=2)
        elif para_text.startswith('[H3]') and '[/H3]' in para_text:
            content_text = para_text.replace('[H3]', '').replace('[/H3]', '')
            doc.add_heading(content_text, level=3)
        elif para_text.startswith('[LI]') and '[/LI]' in para_text:
            content_text = para_text.replace('[LI]', '').replace('[/LI]', '')
            doc.add_paragraph(content_text, style='List Bullet')
        elif para_text.startswith('[CODE]') and '[/CODE]' in para_text:
            code_text = para_text.replace('[CODE]', '').replace('[/CODE]', '')
            p = doc.add_paragraph()
            run = p.add_run(code_text)
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
        elif para_text.startswith('|') and '|' in para_text[1:]:
            # 简单的表格检测
            rows = [row.strip() for row in para_text.split('\n') if '|' in row]
            if len(rows) >= 2:
                cells_first = [c.strip() for c in rows[0].split('|') if c.strip()]
                if cells_first:
                    table = doc.add_table(rows=len(rows), cols=len(cells_first))
                    table.style = 'Light Grid Accent 1'
                    for i, row in enumerate(rows):
                        cells = [c.strip() for c in row.split('|') if c.strip()]
                        for j, cell_text in enumerate(cells):
                            if i < len(table.rows) and j < len(table.rows[i].cells):
                                table.rows[i].cells[j].text = cell_text
                    doc.add_paragraph()  # 表后空行
        else:
            doc.add_paragraph(para_text)


def _lines_to_docx(content: str, doc):
    """逐行转换 Markdown → Word（降级方案）。"""
    from docx.shared import Pt

    for line in content.split("\n"):
        line = line.strip()
        if not line:
            continue
        if line.startswith("# ") and not line.startswith("## "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("```"):
            continue
        elif line.startswith("| "):
            continue  # 表格在 _html_to_docx 中处理
        elif line.startswith("- "):
            doc.add_paragraph(line[2:], style="List Bullet")
        elif line == "---":
            doc.add_paragraph("─" * 40)
        else:
            doc.add_paragraph(line)


def _export_html(content: str, output_dir: Path, filename: str, chart_json: dict | None = None) -> str:
    """导出 HTML — 将 base64 图片转为 <img> 标签再渲染。"""
    import re as _re

    # 先将 ![alt](data:image/png;base64,...) 转为 <img> 标签
    # markdown 库不认识 data: URI，必须提前转换
    content = _re.sub(
        r'!\[([^\]]*)\]\((data:image/png;base64,[^)]+)\)',
        r'<img src="\2" alt="\1" style="max-width:100%;">',
        content,
    )

    try:
        import markdown
        html_body = markdown.markdown(content, extensions=["tables", "fenced_code", "codehilite"])
    except ImportError:
        html_body = content.replace("\n", "<br>\n")

    has_charts = 'data:image/png;base64,' in content or 'chart' in content.lower()

    # 如果有 Plotly JSON，在 HTML body 后追加交互式图表
    chart_div = ""
    if chart_json and isinstance(chart_json, dict) and "data" in chart_json:
        import json
        chart_id = "chart_main"
        chart_json_str = json.dumps(chart_json, ensure_ascii=False)
        chart_div = f"""
<div class="chart-container" style="margin: 24px 0; padding: 16px; background: var(--code-bg); border-radius: 8px; border: 1px solid var(--border);">
    <h3 style="margin-top:0;">📊 交互式图表</h3>
    <div id="{chart_id}" style="width:100%; min-height:450px;"></div>
</div>
<script>
(function() {{
    var data = {chart_json_str};
    if (typeof Plotly !== 'undefined') {{
        Plotly.newPlot('{chart_id}', data.data, data.layout || {{}}, {{responsive: true, displayModeBar: true}});
    }} else {{
        document.getElementById('{chart_id}').innerHTML = '<p style=\"color:#999;padding:40px;text-align:center\">[Plotly.js 未加载 — 图表数据已在页面源码中]</p>';
    }}
}})();
</script>
"""

    html = f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>DataForge AI 分析报告</title>
    {('<script src="https://cdn.plot.ly/plotly-2.32.0.min.js"></script>' if has_charts else '')}
    <style>
        :root {{
            --primary: #409eff;
            --bg: #ffffff;
            --text: #333333;
            --border: #e4e7ed;
            --code-bg: #f8f8f8;
        }}
        @media (prefers-color-scheme: dark) {{
            :root {{
                --bg: #1e1e1e;
                --text: #d4d4d4;
                --border: #404040;
                --code-bg: #2d2d2d;
            }}
        }}
        body {{
            font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", "PingFang SC",
                         "Microsoft YaHei", sans-serif;
            max-width: 900px;
            margin: 0 auto;
            padding: 20px 24px;
            line-height: 1.8;
            color: var(--text);
            background: var(--bg);
        }}
        h1 {{ color: var(--primary); border-bottom: 3px solid var(--primary); padding-bottom: 8px; }}
        h2 {{ color: var(--text); border-bottom: 2px solid var(--border); padding-bottom: 6px; margin-top: 32px; }}
        h3 {{ color: var(--text); margin-top: 24px; }}
        table {{ border-collapse: collapse; width: 100%; margin: 16px 0; }}
        th, td {{ border: 1px solid var(--border); padding: 8px 12px; text-align: left; }}
        th {{ background: var(--primary); color: white; }}
        code {{ background: var(--code-bg); padding: 2px 6px; border-radius: 4px; font-family: "Consolas", "Courier New", monospace; font-size: 0.9em; }}
        pre {{ background: var(--code-bg); padding: 16px; border-radius: 8px; overflow-x: auto; border: 1px solid var(--border); }}
        pre code {{ background: none; padding: 0; }}
        blockquote {{ border-left: 4px solid var(--primary); padding: 8px 16px; margin: 16px 0; color: #666; background: var(--code-bg); }}
        details {{ margin: 12px 0; padding: 12px; border: 1px solid var(--border); border-radius: 8px; }}
        details summary {{ cursor: pointer; font-weight: bold; color: var(--primary); }}
        details[open] {{ background: var(--code-bg); }}
        .chart-container {{ box-shadow: 0 2px 8px rgba(0,0,0,0.08); }}
        img {{ max-width: 100%; height: auto; }}
        a {{ color: var(--primary); }}
    </style>
</head>
<body>
{html_body}
{chart_div}
<hr>
<p><em>本报告由 DataForge AI 自动生成 | 7 Agents + LangGraph | {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</em></p>
</body>
</html>"""

    filepath = output_dir / f"{filename}.html"
    filepath.write_text(html, encoding="utf-8")
    logger.info("HTML 报告已导出: %s (图表=%s)", filepath, has_charts)
    return str(filepath)


# ═══════════════════════════════════════════════════════════
# 工具函数
# ═══════════════════════════════════════════════════════════


def _slugify(text: str) -> str:
    """简单的中文友好 slug。"""
    result = []
    for c in text:
        if c.isalnum() or c in "_-":
            result.append(c)
        elif c.isspace():
            result.append("_")
        else:
            result.append("_")
    return "".join(result).strip("_") or "report"
